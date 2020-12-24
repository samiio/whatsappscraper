"""
This script allows the user to save a WhatsApp conversation from
'web.whatsapp.com' to a (.docx), (.csv), or (.db) file.

This script requires that 'selenium', 'pyinputplus', and 'python-docx'
be installed within the Python environment you are running it in.

Written by: Sami Ali
"""
import re
import csv
import time
import docx
import sqlite3
import pyinputplus as pyip

from sys import exit
from datetime import timedelta
from selenium import webdriver
from selenium.common import exceptions
from selenium.webdriver.common.keys import Keys
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH 

# Regex
dateRegex = re.compile(r"""
    ([0-3][0-9])        # day
    /                   # seperator
    ([0-1][0-9])        # month
    /                   # seperator
    ([1-2][0-9]{3})     # year
    """, re.VERBOSE)

def main():
    """
    Main body of the program
    """
    print("\nPlease follow the instructions carefully:\n"
        + "The WhatsApp website will be launched shortly.\n"
        + "Scan the QR code once the website has loaded.\n"
        + "Return here once you have been logged in.\n"
        + "Press 'Enter' key to continue.")
    input()

    browser = webdriver.Firefox()
    browser.get("https://web.whatsapp.com/")
    print("Press 'Enter' after logging in.")
    input()

    # Prompt user for contact name and valid date
    personName = input("Contact name: ")
    print("\nEnter dates, format: MM/DD/YYYY\n"
          "NOTE: Dates must be listed in the chat")
    while True:
        startDate = pyip.inputDate("From: ")
        endDate = pyip.inputDate("To  : ")
        if startDate < endDate:
            break
        else:
            print("\n'From' is the start date\n")

    print("\nNow return to the browser and wait for the \n"
        + "program to finish execution.\n"
        + "Press 'Enter' to continue.")
    input()
    print("You have 10 seconds to return to the browser")
    countDown(3)

    # Search name and open conversation
    searchBoxElem = browser.find_element_by_class_name("_1Ra05")
    searchBoxElem.click()
    searchBoxElem.send_keys(personName)
    nameElem = findText(browser, personName)
    nameElem.click() if nameElem else exit()

    # Scroll to start and store messages upto end
    scrollToDate(browser, startDate)
    textMessages = getMessages(browser, endDate)
    browser.quit()

    # Prompt user to choose output file format
    print("\nExport messages as:"
          "\n1. Document file"
          "\n2. CSV file"
          "\n3. SQLite file (.db)")
    choice = pyip.inputInt(min=1, max=3)
    if choice == 1:
        sDoc = writeToDoc(personName, textMessages)
        printOutputDetails(
            choice, personName, sDoc, textMessages)
    elif choice == 2:
        sCsv = writeToCsv(personName, textMessages)
        printOutputDetails(
            choice, personName, sCsv, textMessages)
    else:
        sSql = writeToSql(personName, textMessages)
        printOutputDetails(
            choice, personName, sSql, textMessages)


def printOutputDetails(choice, name, size, messages):
    """
    Print name of the saved file and number of messages.

    Args:
        choice (int): 1 for docx, 2 for csv, 3 for db.
        name (string): participent name.
        size (int): num of saved messages from the chat.
        messages (list of tuples): messages in the chat.
    """
    if choice == 1:
        print(f"\n{name}-chat.docx saved.")
    elif choice == 2:
        print(f"\n{name}-chat.csv saved.")
    else:
        print(f"\n{name}-chat.db saved.")

    print(f"{size} / {len(messages)} messages written.")


def writeToDoc(name, messages):
    """
    Write WhatsApp messages from conversation with 'name' 
    into document file.

    Args:
        name (string): participent name.
        messages (list of tuples): messages in the chat.

    Returns: 
        int: number of messages written to file.
    """
    doc = docx.Document()
    for msg in messages:
        pInfo = doc.add_paragraph()
        pMsg = doc.add_paragraph(msg[1])
        rInfo = pInfo.add_run(msg[0])
        rInfo.bold = True
        font = rInfo.font

        # Use different font color for memebers in chat
        if msg[0].find(name) != -1:
            font.color.rgb = RGBColor(0x80, 0x00, 0x20)
        else:
            font.color.rgb = RGBColor(0x00, 0x00, 0x80)
            pInfo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pMsg.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Paragraph formatting
        pFormat = doc.styles['Normal'].paragraph_format
        pFormat.space_after = 0
        paragraph = doc.add_paragraph()
    doc.save(f"{name}-chat.docx")
    return int(len(doc.paragraphs) / 3)


def writeToCsv(name, messages):
    """
    Write WhatsApp messages from conversation with 'name' 
    into csv file.

    Args:
        name (string): participent name.
        messages (list of tuples): messages in the chat.

    Returns: 
        int: number of messages written to file.
    """
    with open(f"{name}-chat.csv", "w", newline='') as f:
        writer = csv.writer(f)
        writer.writerow(["info", "message"])
        for msg in messages:
            writer.writerow(list(msg))
    
    # Open file to count rows
    with open(f"{name}-chat.csv", "r") as f:
        reader = csv.reader(f, delimiter = ",")
        data = list(reader)
        return len(data) - 1


def writeToSql(name, messages):
    """
    Insert WhatsApp messages from conversation with 'name' 
    into SQLite database.

    Args:
        name (string): participent name.
        messages (list of tuples): messages in the chat.

    Returns: 
        int: number of messages written to database.
    """
    open(f"{name}-chat.db", "w")
    db = sqlite3.connect(f"{name}-chat.db")
    db.execute("CREATE TABLE chat (info TEXT, message TEXT)")
    for msg in messages:
        c = db.cursor()
        c.execute(
            "INSERT INTO chat (info, message) VALUES (?, ?)",
            msg)
        db.commit()
        c.close()
    
    # Execute query to count rows
    c = db.cursor()
    c.execute("SELECT COUNT(*) FROM chat")
    nRows = c.fetchone()[0]
    c.close()
    return nRows


def getMessages(browser, eDate):
    """
    Get WhatApp messages, text only, up to a certain date.

    Args:
        browser (selenium.webdriver): Firefox driver.
        eDate (datetime.date): date to stop at.

    Returns: 
        list: tuples of (time + sender)-message pairs.
    """
    messages = []
    eDateStr = str(eDate.strftime("%d/%m/%Y"))
    msgElem = browser.find_elements_by_class_name("copyable-text")
    for msg in msgElem:
        try:
            if msg.text:
                senderInfo = msg.get_attribute("data-pre-plain-text")
                if senderInfo and senderInfo.find(eDateStr) != -1:
                    break
                if senderInfo:
                    messages.append((senderInfo, msg.text))
        except exceptions.StaleElementReferenceException:
            scrollElem = browser.find_element_by_class_name("_26MUt")
            scrollElem.click()
            scrollElem.send_keys(Keys.DOWN)
    return messages


def scrollToDate(browser, sDate):
    """
    Scroll to specified date in the chat open in WhatsApp Web.

    Args:
        browser (selenium.webdriver): Firefox driver.
        sDate (datetime.date): date to scroll to date.
    """
    sDateN = str(
        (sDate + timedelta(days=-1)).strftime("%d/%m/%Y"))
    scrollElem = browser.find_element_by_class_name("_26MUt")
    scrollElem.click()
    while findText(browser, sDateN) is None:
        scrollElem.send_keys(Keys.UP)
    

def findText(browser, key):
    """
    Search for text key in the webpage.

    Args:
        browser (selenium.webdriver): Firefox driver.
        key (string): text to find.

    Returns: 
        WebElement: of the key if it exists, None otherwise.
    """
    try:
        return browser.find_element_by_xpath(
            "//*[contains(text(), '" + key + "')]")
    except exceptions.NoSuchElementException:
        if re.search(dateRegex, key) == None:
            print(key, "not found.")
    return None


def countDown(n):
    """
    Countdown from 'n' to 0.

    Args:
        n (int): the start
    """
    for i in range(n, 0, -1):
        print(str(i))
        time.sleep(1)
    

if __name__ == "__main__":
    main()
